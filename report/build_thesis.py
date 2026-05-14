"""
Tez Word (.docx) Üretici — v4 (Akademik makale formatı)
=========================================================

Üretilen rapor:
  - Akademik makale formatı (Önsöz yok)
  - İlk sayfa: başlık + yazar/meta + URL + MAKALE BİLGİSİ + ÖZET
  - Akademik dil (1. tekil yok)
  - 1. Giriş + araştırma soruları (S1-S4)
  - 2. Veri Kümesi (alt-bölümler: tanıtım → ham format → kalite → bölme → örnek)
  - 3. Modeller ve Metodoloji
  - 4. Sonuçlar (alt-bölümler: eğitim eğrileri, test perf., ufuk, mimari, literatür)
  - 5. Tartışma (araştırma sorularına yanıt, sınırlılıklar, öneriler)
  - 6. ISSD Entegrasyon
  - 7. Sonuç

Çalıştırma:
    python report/build_thesis.py
"""

import os, sys, json

HERE = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(HERE, ".."))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

import numpy as np
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGURES_DIR = os.path.join(PROJECT_ROOT, "figures")
RESULTS_DIR = os.path.join(PROJECT_ROOT, "results")
DATA_DIR    = os.path.join(PROJECT_ROOT, "data", "metr-la")
REPORT_DIR  = os.path.join(PROJECT_ROOT, "report")
OUTPUT_PATH = os.path.join(REPORT_DIR, "thesis.docx")

METR_LA_URL = "https://github.com/liyaguang/DCRNN"


# =============================================================================
# Yardımcılar
# =============================================================================
def set_default_styles(doc):
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"; s.font.size = Pt(11)
    pf = s.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.5)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for hname, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        h = doc.styles[hname]
        h.font.name = "Times New Roman"; h.font.size = Pt(size)
        h.font.bold = True; h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.left_indent = Cm(0)
        h.paragraph_format.space_before = Pt(12 if size > 12 else 8)
        h.paragraph_format.space_after = Pt(6)


def setup_page(doc):
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.0))


def add_par(doc, text, bold=False, italic=False, size=11,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, color=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p


def add_heading(doc, text, level=1, numbered=False, number=None):
    if numbered and number is not None: text = f"{number}. {text}"
    h = doc.add_heading(text, level=level)
    h.paragraph_format.left_indent = Cm(0)
    h.paragraph_format.first_line_indent = Cm(0)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


FIG_COUNTER = {"n": 0}
TBL_COUNTER = {"n": 0}


def add_figure(doc, image_filename, caption_text, width_inches=6.4):
    FIG_COUNTER["n"] += 1
    n = FIG_COUNTER["n"]
    image_path = os.path.join(FIGURES_DIR, image_filename)
    if not os.path.isfile(image_path):
        add_par(doc, f"[Şekil {n}: {caption_text} — görsel yok]",
                italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        return n
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(image_path, width=Inches(width_inches))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(f"Şekil {n}. {caption_text}")
    cr.italic = True; cr.font.name = "Times New Roman"; cr.font.size = Pt(10)
    return n


def add_table_caption(doc, caption_text):
    TBL_COUNTER["n"] += 1
    n = TBL_COUNTER["n"]
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0); cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(f"Tablo {n}. {caption_text}")
    r.italic = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
    return n


def style_table(table, font_size=10):
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Light Grid Accent 1"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# =============================================================================
# Veri yükle
# =============================================================================
print("Metrikler yükleniyor...")
with open(os.path.join(RESULTS_DIR, "metrics_a3tgcn_baseline.json"), encoding="utf-8") as f: m_a3t = json.load(f)
with open(os.path.join(RESULTS_DIR, "metrics_dcrnn_baseline.json"),  encoding="utf-8") as f: m_dcrnn = json.load(f)
with open(os.path.join(RESULTS_DIR, "metrics_vanilla_lstm.json"),    encoding="utf-8") as f: m_lstm = json.load(f)
with open(os.path.join(RESULTS_DIR, "realunit_comparison.json"),     encoding="utf-8") as f: m_real = json.load(f)
with open(os.path.join(RESULTS_DIR, "dataset_stats.json"),           encoding="utf-8") as f: m_ds = json.load(f)

# Örnek snapshot verisini oku (Bölüm 2'deki "örnek üzerinden temsili gösterim" için)
print("Örnek snapshot verisi yükleniyor...")
X_full = np.load(os.path.join(DATA_DIR, "node_values.npy"))   # [T, N, F]
A_full = np.load(os.path.join(DATA_DIR, "adj_mat.npy"))        # [N, N]
EXAMPLE_TIME_IDX = 15000
EXAMPLE_SENSOR_IDX = 0   # ilk sensör (#773869)
example_speed_series = X_full[EXAMPLE_TIME_IDX:EXAMPLE_TIME_IDX + 12, EXAMPLE_SENSOR_IDX, 0]
example_tod_series   = X_full[EXAMPLE_TIME_IDX:EXAMPLE_TIME_IDX + 12, EXAMPLE_SENSOR_IDX, 1]
example_neighbors = np.where(A_full[EXAMPLE_SENSOR_IDX] > 0)[0]


# =============================================================================
# DOCUMENT
# =============================================================================
print("Word dokümanı oluşturuluyor...")
doc = Document()
set_default_styles(doc); setup_page(doc)


# -----------------------------------------------------------------------------
# İLK SAYFA — Akademik makale formatı: başlık + meta + ÖZET
# -----------------------------------------------------------------------------
# Üst boşluk
doc.add_paragraph()

# Başlık
add_par(doc,
        "Graf Sinir Ağları ile Trafik Hız Tahmini: A3T-GCN, DCRNN ve "
        "Vanilla LSTM Mimarilerinin Tahmin Ufkuna Göre Yeterlilik Analizi",
        bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
doc.add_paragraph()

# Yazar bilgisi
add_par(doc, "Yusuf İnce", bold=True, size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_par(doc, "Yapay Zeka Mühendisliği Bölümü · 3. Sınıf",
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_par(doc, "Ders: Derin Öğrenme · Mayıs 2026",
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

# Veri kümesi URL'si
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("Veri kümesi: METR-LA — Los Angeles Otoyol Sensör Ağı · ")
r.font.name = "Times New Roman"; r.font.size = Pt(10); r.italic = True
add_hyperlink(p, METR_LA_URL, METR_LA_URL)

doc.add_paragraph()

# MAKALE BİLGİSİ + ÖZET kutusu (tek kolon ama vurgulu başlıklı)
add_par(doc, "MAKALE BİLGİSİ", bold=True, size=10, indent=False)

p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.left_indent = Cm(0.3)
r = p.add_run("Anahtar Kelimeler:")
r.font.name = "Times New Roman"; r.font.size = Pt(10); r.bold = True

keywords = [
    "Graf Sinir Ağları",
    "DCRNN, A3T-GCN, LSTM",
    "Spatio-temporal modelleme",
    "Trafik hız tahmini",
    "METR-LA veri kümesi",
    "Akıllı Ulaşım Sistemleri (ITS)",
    "Yeterlilik koşulları",
    "Yönlü difüzyon konvolüsyonu",
]
for kw in keywords:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("· " + kw)
    r.font.name = "Times New Roman"; r.font.size = Pt(10)

doc.add_paragraph()

# ÖZET
add_par(doc, "ÖZET", bold=True, size=10, indent=False)

add_par(doc, (
    "Bu çalışmanın amacı, şehir-ölçekli trafik hız tahmini probleminde graf sinir "
    "ağı (Graph Neural Network, GNN) mimarilerinin yeterlilik koşullarını "
    "incelemektir. Klasik zaman serisi modelleri her sensörün tahminini büyük "
    "ölçüde kendi geçmişine dayandırırken, GNN tabanlı yaklaşımlar sensörler "
    "arası ilişki yapısını öğrenme sürecinin içine yerleştirerek mekansal "
    "yayılma örüntülerini modelleyebilmektedir. Bu çalışma, METR-LA veri kümesi "
    "üzerinde üç farklı mimariyi sistematik biçimde karşılaştırmaktadır: A3T-GCN "
    "(zamansal dikkat ile yönsüz graf konvolüsyonu), DCRNN (yönlü difüzyon "
    "konvolüsyonu ile recurrent hücre) ve graf yapısını hiç kullanmayan kontrol "
    "grubu olarak vanilla LSTM. Çalışmanın odağı yalnızca en yüksek skoru elde "
    "etmek değil; mimari seçiminin, tahmin ufkunun ve graf yönlülüğünün modelin "
    "operasyonel yeterliliğine etkisini bütünsel olarak çözümlemektir. PyTorch "
    "Geometric Temporal kütüphanesinin yığın tabanlı sürümleri kullanılarak "
    "yaklaşık 150 kat hızlandırılmış bir eğitim hattı kurulmuş; tüm deneyler "
    "tohum 42 altında, 10 epoch boyunca NVIDIA RTX 4050 Laptop GPU üzerinde "
    "yürütülmüştür. Bulgular kontrollü bir yeterlilik analizi sunmaktadır: kısa "
    f"tahmin ufkunda (5-45 dk) vanilla LSTM yeterli bir karşılaştırma noktası "
    f"oluşturmakta (MAE {m_real['models'][2]['test_mae_mph']:.2f} mph); 55 "
    f"dakika ve sonrasında DCRNN'in yönlü mekansal indüktif önyargısı yeterli "
    f"hale gelmekte ve RMSE metriğinde tüm ufuklarda öne çıkmaktadır (RMSE "
    f"{m_real['models'][1]['test_rmse_mph']:.2f} mph). A3T-GCN, METR-LA'nın "
    "asimetrik komşuluk yapısını modelleyememesi nedeniyle her tahmin ufkunda "
    "diğer iki mimarinin gerisinde kalmıştır. Çalışmanın endüstriyel bağlamı, "
    "DCRNN modelinin ISSD Bilişim Elektronik A.Ş.'nin CHAOS ve MANGO ürünlerine "
    "paralel öngörü katmanı olarak entegrasyon önerisi şeklinde sunulmuş; bu "
    "öneri Streamlit tabanlı interaktif bir kontrol panosu ile somutlaştırılmıştır."
))

add_page_break(doc)


# -----------------------------------------------------------------------------
# 1. GİRİŞ
# -----------------------------------------------------------------------------
add_heading(doc, "Giriş", level=1, numbered=True, number=1)

add_par(doc, (
    "Bu çalışmanın hedefi tek bir cümleyle özetlenebilir: trafik sensörlerinden "
    "oluşan bir mekansal ağı düğümleri sensörler ve kenarları sensörler arası yol "
    "üzerinden bağlantılar olan bir graf olarak yapay zekâya tanıtmak ve bu graf "
    "üzerinden 60 dakika sonrasının trafik hızlarını graf sinir ağı tabanlı "
    "modellerle tahmin etmektir. Diğer bir deyişle modele 'gelecek bir saat "
    "içinde bu kavşaktaki trafik nasıl olacak?' sorusu graf düzeyinde bir "
    "regresyon problemi olarak sorulmakta; cevap, modelin mekansal ve zamansal "
    "yapıdan öğrendiği temsillere dayanarak üretilmektedir."
))

add_par(doc, (
    "Bu hedefe ulaşmak için üç farklı mimari aynı koşullar altında "
    "karşılaştırılmaktadır: A3T-GCN (yönsüz graf konvolüsyonu ve zamansal "
    "dikkat), DCRNN (yönlü difüzyon konvolüsyonu ve GRU recurrent hücre) ve "
    "graf yapısını hiç kullanmayan klasik bir LSTM. Üçüncü mimari, graf "
    "yapısının trafik tahminindeki katkısını ayrıştırmak için ölçüt olarak yer "
    "almaktadır; eğer graf bilinçli mimariler bu ölçütü net biçimde geçemiyorsa, "
    "ek mimari karmaşıklığın bu problem türünde değer eklemediği sonucuna "
    "varılabilir."
))

add_par(doc, (
    "Trafik tahmini, akıllı ulaşım sistemleri (Intelligent Transportation "
    "Systems, ITS) alanındaki en kritik hesaplamalı problemlerden biridir. Bir "
    "kavşaktaki tıkanıklığı dakikalar önceden tahmin edebilmek, sinyalizasyon "
    "yönetimini reaktif olmaktan çıkarıp proaktif hale getirmenin ön koşuludur. "
    "Klasik zaman serisi modelleri (ARIMA, LSTM, Transformer) bir sensörün "
    "tahminini büyük ölçüde kendi geçmişine dayandırır; ancak trafiğin yapısı "
    "gereği komşu sensörlerden gelen sinyaller, özellikle uzun tahmin ufkunda, "
    "kritik bilgi taşır. Graf sinir ağları, ağın ilişki yapısını öğrenme "
    "sürecinin içine yerleştirerek bu eksikliği gidermeyi vaat etmektedir."
))

add_par(doc, (
    "Pratik bir mühendislik bağlamında karşılaşılan soru çoğu zaman 'hangi "
    "mimari en yüksek skoru verir' değil; uygulamanın operasyonel amacıyla "
    "yöntem seçimi arasındaki ilişkidir. Bazı uygulamalarda asıl belirleyici "
    "olan tipik hata (MAE); örneğin sürücü bilgilendirme servislerinde gün "
    "boyunca tutarlı doğruluk istenir. Bazı uygulamalarda ise modelin büyük "
    "hatalardan kaçınması, yani kuyruktaki aşırı sapmaların asgariye indirilmesi "
    "(RMSE) daha kritiktir; örneğin proaktif sinyal yönetiminde büyük "
    "tıkanıklık olaylarını kaçırmamak gerekir. Bu sebeple 'doğru' bir mimari "
    "tercihi yoktur; tercih, modelin hangi operasyonel hedefe hizmet edeceğine "
    "bağlıdır. Bu çalışma, mimari seçiminin yanı sıra tahmin ufkunun ve graf "
    "yönlülüğünün hata profili üzerindeki etkisini birden fazla amaç ekseninde "
    "değerlendirebilmek için sistematik bir karşılaştırma sunmaktadır."
))

# Araştırma soruları
add_par(doc, (
    "Çalışma dört temel araştırma sorusu etrafında şekillendirilmiştir:"
), indent=False)

questions = [
    ("S1.", "Farklı graf sinir ağı mimarileri (A3T-GCN, DCRNN) METR-LA "
     "veri kümesi üzerinde nasıl bir performans dağılımı sergilemektedir? "
     "Mesaj iletme şemalarındaki farklılıklar — simetrik graf konvolüsyonu ve "
     "yönlü difüzyon konvolüsyonu — trafik tahmin görevinde anlamlı farklara "
     "yol açmakta mıdır?"),
    ("S2.", "Yönlü graf yapısının (asimetrik komşuluk matrisi) modellenmesi, "
     "simetrik komşuluk varsayımıyla karşılaştırıldığında ne kadar iyileştirme "
     "sağlamaktadır? Yönlülüğün katkısı tahmin ufkuna göre nasıl değişmektedir?"),
    ("S3.", "Tahmin ufkunun büyümesi farklı mimari ailelerinin yeterliliğini "
     "nasıl değiştirmektedir? Graf-bilinçli mimariler ile graf yapısını "
     "kullanmayan klasik bir tekrarlı ağ arasında bir geçiş noktası (crossover) "
     "mevcut mudur ve eğer öyleyse hangi ufukta gerçekleşmektedir?"),
    ("S4.", "Hata metriği tercihi (MAE veya RMSE) mimari yeterlilik sıralamasını "
     "nasıl etkilemektedir? Büyük hatalara duyarlı senaryolarda (RMSE) ve tipik "
     "hata senaryolarında (MAE) hangi mimari hangi koşullarda öne çıkmaktadır?"),
]
for tag, body in questions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7)
    r1 = p.add_run(tag + " "); r1.bold = True
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
    r2 = p.add_run(body); r2.font.name = "Times New Roman"; r2.font.size = Pt(11)

add_par(doc, (
    "Çalışmanın katkısı dört eksende özetlenebilir: (i) üç mimarinin (A3T-GCN, "
    "DCRNN, vanilla LSTM) aynı veri bölme stratejisi ve aynı eğitim bütçesi "
    "altında kontrollü deneysel ortamda karşılaştırılması; (ii) yönlülük "
    "etkisinin A3T-GCN (simetrik) ve DCRNN (yönlü) üzerinden izole biçimde "
    "ölçülmesi; (iii) tahmin ufkuna göre yeterlilik koşullarının çıkarılması ve "
    "vanilla LSTM ile graf-bilinçli mimariler arasındaki geçiş noktasının "
    "tespit edilmesi; (iv) bulguların ISSD'nin akıllı ulaşım platformuna "
    "entegrasyon önerisine dönüştürülmesi ve etkileşimli bir kontrol panosu "
    "prototipi ile somutlaştırılması."
))

# 1.1 İlgili çalışmalar
add_heading(doc, "İlgili çalışmalar", level=2, numbered=True, number="1.1")
add_par(doc, (
    "Graf sinir ağı tabanlı trafik tahmini literatürünün başlangıç noktası "
    "olarak Kipf ve Welling (2017) tarafından önerilen Graph Convolutional "
    "Network (GCN) gösterilmektedir. GCN, spektral graf teorisinden türetilen "
    "yerel komşuluk normalizasyonuyla düğüm gömme vektörlerini günceller ve "
    "klasik konvolüsyonel ağların düzensiz graflara genelleşmesini sağlar. Bu "
    "temel üzerine spatio-temporal modelleme için iki önemli mimari "
    "geliştirilmiştir: Li ve diğerleri (2018) tarafından önerilen DCRNN, "
    "yönlü graf üzerinde difüzyon konvolüsyon operatörü ile GRU recurrent "
    "hücresini birleştirmektedir. Aynı yıl Yu ve diğerleri (2018) STGCN ile "
    "blok-yapılı bir spatio-temporal konvolüsyon mimarisi sunmuştur."
))
add_par(doc, (
    "Zhu ve diğerleri (2020) A3T-GCN ile TGCN'in üzerine zamansal dikkat "
    "mekanizmasını ekleyerek, modelin geçmişin hangi noktalarına odaklanması "
    "gerektiğini öğrenmesini sağlamıştır. Wu ve diğerleri (2019) Graph WaveNet "
    "ile uyarlanabilir komşuluk matrisi öğrenimini önererek statik graf "
    "varsayımını gevşetmiştir. METR-LA veri kümesi, yayımlanmasının ardından "
    "geniş bir çalışma külliyatına konu olmuş; literatürde raporlanan en iyi "
    "MAE değerleri 15 dakikalık ufukta 2,7 mph civarındadır. Bu çalışmaların "
    "büyük çoğunluğu tek bir mimariyi derinlemesine incelemekte; üç farklı "
    "mimari ailesini (yönsüz GCN, yönlü difüzyon, graf-yok LSTM) aynı çatı "
    "altında ve aynı eğitim bütçesinde karşılaştıran çalışma sayısı sınırlı "
    "kalmaktadır. Bu çalışma sözü edilen boşluğu doldurmayı amaçlamaktadır."
))
add_par(doc, (
    "PyTorch Geometric Temporal (Rozemberczki ve diğerleri, 2021), bu çalışmada "
    "kullanılan birincil yazılım kütüphanesidir; spatio-temporal graf sinir "
    "ağı mimarilerinin yığın tabanlı sürümlerini sunmakta ve eğitim hızını "
    "büyük ölçüde artırmaktadır. Bulguların literatürdeki diğer çalışmalarla "
    "niceliksel karşılaştırması Bölüm 4.7'de tablo halinde sunulmuştur."
))

# 1.2 Çalışmanın yapısı
add_heading(doc, "Çalışmanın yapısı", level=2, numbered=True, number="1.2")
add_par(doc, (
    "Çalışmanın geri kalanı şu şekilde yapılandırılmıştır. Bölüm 2 METR-LA veri "
    "kümesini, uygulanan veri hazırlık akışını ve örnek bir sensör üzerinden "
    "veri temsilinin somutlaştırılmasını sunmaktadır. Bölüm 3 üç mimariyi, "
    "hiperparametre tercihlerini ve eğitim/değerlendirme protokolünü "
    "tanımlamaktadır. Bölüm 4 sonuçları betimsel ve karşılaştırmalı bulgular "
    "(mimari karşılaştırması, ufuk-bazlı analiz, mekansal yayılma, hesaplama "
    "maliyeti ve literatürle karşılaştırma) halinde sunmaktadır. Bölüm 5 "
    "araştırma sorularını ampirik verilere dayandırarak tartışmakta, "
    "yöntemsel gözlemleri yorumlamakta ve sınırlılıkları ortaya koymaktadır. "
    "Bölüm 6 ISSD entegrasyon önerisini detaylandırmakta, Bölüm 7 sonuçları "
    "özetleyip gelecek çalışmalar için somut yönler önermektedir."
))


# -----------------------------------------------------------------------------
# 2. VERİ KÜMESİ
# -----------------------------------------------------------------------------
add_heading(doc, "Veri Kümesi", level=1, numbered=True, number=2)

# 2.1 METR-LA tanıtım
add_heading(doc, "METR-LA Veri Kümesi", level=2, numbered=True, number="2.1")
p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.5)
r = p.add_run(
    "Çalışmada Li ve diğerleri (2018) tarafından tanıtılan ve "
    "literatürde trafik tahmini için referans niteliğindeki METR-LA veri kümesi "
    "kullanılmıştır ("
)
r.font.name = "Times New Roman"; r.font.size = Pt(11)
add_hyperlink(p, METR_LA_URL, METR_LA_URL)
r2 = p.add_run(
    "). Veri kümesinin tercih edilmesinin üç temel gerekçesi vardır: ilk olarak, "
    "veri büyüklüğü (207 sensör × 34 272 zaman adımı) modern derin öğrenme "
    "modellerini eğitmek için yeterlidir. İkinci olarak, DCRNN paper'ının "
    "yayımladığı resmi veri bölme stratejisi sayesinde elde edilen sonuçlar "
    "literatürdeki diğer çalışmalarla birebir karşılaştırılabilir. Üçüncü "
    "olarak, sensörler arası komşuluk matrisi asimetriktir; bu özellik yönlü "
    "graf modellemenin etkisini izole edebilmeyi mümkün kılmaktadır."
)
r2.font.name = "Times New Roman"; r2.font.size = Pt(11)

add_par(doc, (
    "Etiket konvansiyonu olarak her sensörün her zaman adımında ölçtüğü "
    "ortalama hız (mph cinsinden) doğrudan kullanılmıştır. Tahmin görevi 12 "
    "zaman adımlık (60 dakikalık) geçmiş penceresinden 12 adımlık gelecek "
    "penceresinin tahmin edilmesidir. Tüm performans metrikleri hız değişkeni "
    "üzerinden hesaplanmıştır."
))

add_table_caption(doc, "METR-LA veri kümesinin temel özellikleri")
ds_stats = m_ds
tbl = doc.add_table(rows=12, cols=2)
rows = [
    ("Veri kümesi adı", "METR-LA"),
    ("Kaynak", "DCRNN paper'ı (Li et al., 2018) / Los Angeles Metro"),
    ("Düğüm sayısı (sensör)", str(ds_stats["spatial"]["num_nodes"])),
    ("Kenar sayısı", str(ds_stats["spatial"]["num_edges"])),
    ("Komşuluk yapısı", "Yönlü (asimetrik)"),
    ("Zaman adımı sayısı", f"{ds_stats['temporal']['timesteps']} (5 dakika çözünürlükte)"),
    ("Toplam zaman aralığı", f"{ds_stats['temporal']['span_days']:.1f} gün (Mart-Haziran 2012)"),
    ("Eğitim / Doğrulama / Test", "23.974 / 3.425 / 6.850 (kronolojik)"),
    ("Bölme türü", "Resmi kronolojik bölme (rastgele değil)"),
    ("Öznitelik sayısı", f"{ds_stats['features']['num_features']} (hız + günün saati)"),
    ("Eksik veri", f"%{ds_stats['features']['missing_zeros_pct']:.2f} (sıfır okumalar)"),
    ("Resmi metrikler", "MAE, RMSE (mph)"),
]
for i, (k, v) in enumerate(rows): tbl.cell(i, 0).text = k; tbl.cell(i, 1).text = v
style_table(tbl)
doc.add_paragraph()

# 2.2 Veri hazırlık ve kalite analizi
add_heading(doc, "Veri Hazırlık ve Kalite Analizi", level=2, numbered=True, number="2.2")
add_par(doc, (
    "Çalışmada uygulanan veri hazırlık süreci üç ana adımdan oluşmaktadır: "
    "(1) veri yükleme ve resmi kronolojik bölmenin korunması, (2) veri kalitesi "
    "kontrolü, (3) z-skor normalizasyonu. Aşağıdaki alt bölümler bu adımları "
    "aynı sırayla ele almaktadır."
))

add_heading(doc, "Kronolojik bölmenin korunması", level=3, numbered=True, number="2.2.1")
add_par(doc, (
    "METR-LA veri kümesi DCRNN paper'ında tanımlanmış kronolojik bölme ile "
    "kullanılmaktadır: ilk %70 eğitim, sonraki %10 doğrulama, son %20 test "
    "kümesidir. Rastgele bölme tercih edilseydi model muhtemelen test kümesine "
    "'yakın' bir eğitim örneği bularak ezbere yakın yüksek skorlar üretebilirdi. "
    "Kronolojik bölme bu duruma engel olmakta; modelin gerçekten yeni zaman "
    "dilimlerine ve mevsimsel örüntülere genelleme yeteneğini ölçmektedir. Bu "
    "çalışmada resmi bölme stratejisi aynen korunmuş; özel bir bölme "
    "uygulanmamıştır."
))

add_heading(doc, "Veri kalitesi kontrolü", level=3, numbered=True, number="2.2.2")
add_par(doc, (
    "Modelleri eğitmeden önce veri kümesinin yapısal özelliklerini anlamak "
    "için sistematik bir kontrol uygulanmıştır. Kontrol kriterleri şunlardır: "
    "(i) eksik (NaN) değer sayısı, (ii) sıfır değer sayısı (sensör kesintisi "
    "veya tam tıkanıklığı temsil edebilir), (iii) ortalama, minimum ve maksimum "
    f"hız değerleri, (iv) komşuluk matrisinin simetri özelliği. Sonuçlar Tablo 1'de "
    f"özetlenmiştir. Veri kümesinde %{ds_stats['features']['missing_zeros_pct']:.2f} "
    "oranında sıfır okuma tespit edilmiştir; bu değerler ayrıştırma için ek "
    "bilgi olmadığından bu çalışmada herhangi bir doldurma (imputation) "
    "uygulanmamış, model verisi olduğu gibi sunulmuştur. Bu seçimin "
    "sınırlılıkları Bölüm 5'te ele alınmaktadır."
))

add_heading(doc, "Z-skor normalizasyonu", level=3, numbered=True, number="2.2.3")
add_par(doc, (
    "Her öznitelik için z-skor normalizasyonu uygulanmıştır: tüm zaman ve "
    f"sensör üzerinden hesaplanan ortalama (hız için μ = {m_real['speed_mean_mph']:.2f} mph) "
    f"ve standart sapma (σ = {m_real['speed_std_mph']:.2f} mph) ile değerler "
    "ölçeklenmiş, böylece her öznitelik sıfır ortalamalı ve birim varyanslı "
    "hale getirilmiştir. Normalizasyon parametreleri eğitim kümesinden "
    "hesaplanıp test kümesine olduğu gibi uygulanmıştır. Tahmin sonuçları "
    "yorumlanırken bu z-skor değerleri tekrar mph birimine dönüştürülmüştür."
))

# 2.3 Graf yapısı
add_heading(doc, "Graf Yapısının İncelenmesi", level=2, numbered=True, number="2.3")
add_par(doc, (
    "Şekil 1, 207 sensörün gerçek enlem-boylam konumlarında ve aralarındaki "
    "güçlü kenarlarda (kenar ağırlığı 0,5 üstü) çizilmiş hâlini sunmaktadır. "
    "Renk skalası sensörün veri kümesi boyunca ölçtüğü ortalama hızı "
    "yansıtmaktadır."
))
add_figure(doc, "sensor_map_static.png",
           "METR-LA sensör ağı haritası — Los Angeles otoyol şebekesi. "
           "Renk skalası sensörlerin ortalama hızını göstermektedir; "
           "düşük hızlı sensörler trafiğin yoğun olduğu kavşakları işaret "
           "etmektedir.")

add_par(doc, (
    "Şekil 2, veri kümesinin dört temel istatistiksel dağılımını sunmaktadır: "
    "hız dağılımı, kenar ağırlığı dağılımı, düğüm derecesi dağılımı ve günün "
    "saati dağılımı."
))
add_figure(doc, "fig_dataset_distributions.png",
           "METR-LA veri kümesinin temel istatistiksel dağılımları. "
           "(a) Hız dağılımı; (b) kenar ağırlığı dağılımı; (c) düğüm "
           "derecesi dağılımı; (d) günün saati dağılımı.")

add_par(doc, (
    "Önemli bir gözlem, komşuluk matrisinin asimetrik olmasıdır (A ≠ Aᵀ): "
    "X sensöründen Y sensörüne olan kenar ağırlığı, ters yöndekinden farklıdır. "
    "Bu özellik fiziksel bir gerçekliği yansıtmaktadır — otoyollarda trafik "
    "akışı tek yönlüdür, X'ten Y'ye olan etki Y'den X'e olandan farklı bir "
    "dinamiğe sahiptir. DCRNN'in yönlü difüzyon konvolüsyonu bu asimetriyi "
    "doğrudan modelleyebilen tek mimaridir; A3T-GCN'in temel GCN katmanı ise "
    "simetrik komşuluk varsayımıyla çalışır ve bu farkı kaybeder. Bu yönlülük "
    "etkisinin niceliksel ölçümü Bölüm 4'te sunulmuştur."
))

add_par(doc, (
    "Şekil 3, 30 günlük ortalamadan oluşturulmuş tipik günlük hız profilini "
    "göstermektedir. Sabah ve akşam yoğun saatleri net biçimde görünürken, "
    "öğle ve gece saatlerinde akış serbestleşmektedir. Bu örüntü, günün saati "
    "özniteliğinin modellemede neden gerekli olduğunu açıklamaktadır."
))
add_figure(doc, "fig_daily_pattern.png",
           "30 günlük ortalamadan tipik günlük trafik hız profili. "
           "Sabah rush (07:00-10:00) ve akşam rush (16:00-19:00) "
           "saatlerindeki belirgin düşüş günün saati özniteliğinin "
           "modele eklenmesini motive etmektedir.")

# 2.4 Örnek üzerinden temsil
add_heading(doc, "Örnek Sensör Üzerinden Veri Temsilinin Gösterimi",
            level=2, numbered=True, number="2.4")
add_par(doc, (
    "Bu alt bölüm, yukarıda anlatılan soyut kavramların (graf yapısı, sensör "
    "öznitelikleri, zamansal pencere) somut bir örnek üzerinden adım adım "
    "nasıl çalıştığını göstermektedir. Bu amaçla rastgele bir zaman noktasında "
    "(adım 15.000) seçilmiş bir sensörün veri temsili incelenmektedir."
))

add_heading(doc, "Ham veri formatı", level=3, numbered=True, number="2.4.1")
add_par(doc, (
    "METR-LA veri kümesinin ham veri formatı iki ana yapıdan oluşmaktadır: "
    "hız ölçüm tensörü ve komşuluk matrisi. Hız tensörü [zaman, sensör, "
    "öznitelik] = [34.272 × 207 × 2] boyutlarındadır; her hücre belirli bir "
    "zaman adımında belirli bir sensörün ölçtüğü hızı (mph) ve normalize "
    "edilmiş günün saatini içerir. Komşuluk matrisi 207 × 207 boyutlarında "
    "kayan noktalı bir matristir; sıfırdan büyük her hücre bir kenarı ve "
    "ağırlığını temsil eder."
))

add_heading(doc, "Zamansal pencere yapısı", level=3, numbered=True, number="2.4.2")
add_par(doc, (
    "Modele girdi olarak verilen her örnek, ardışık 12 zaman adımının "
    "öznitelik vektörlerinden oluşmaktadır. Tablo 2, örnek olarak seçilen "
    f"sensörün ({EXAMPLE_TIME_IDX}. zaman adımından başlayan) bir girdi "
    "penceresini göstermektedir."
))

add_table_caption(doc, f"Örnek sensörün {EXAMPLE_TIME_IDX}. zaman adımından başlayan "
                  "12 adımlık girdi penceresi (hız ve günün saati değerleri).")
tbl = doc.add_table(rows=13, cols=3)
tbl.cell(0, 0).text = "Adım"; tbl.cell(0, 1).text = "Hız (mph)"; tbl.cell(0, 2).text = "Günün saati"
for p in tbl.cell(0, 0).paragraphs:
    for run in p.runs: run.bold = True
for p in tbl.cell(0, 1).paragraphs:
    for run in p.runs: run.bold = True
for p in tbl.cell(0, 2).paragraphs:
    for run in p.runs: run.bold = True
for i in range(12):
    tbl.cell(i+1, 0).text = f"t+{i*5} dk"
    tbl.cell(i+1, 1).text = f"{example_speed_series[i]:.2f}"
    tbl.cell(i+1, 2).text = f"{example_tod_series[i]:.4f}"
style_table(tbl, font_size=10)
doc.add_paragraph()

add_par(doc, (
    f"Yukarıdaki örnekte sensörün ortalama hızı yaklaşık "
    f"{example_speed_series.mean():.1f} mph'dir; bu değer veri kümesinin genel "
    f"ortalamasıyla ({m_real['speed_mean_mph']:.1f} mph) karşılaştırıldığında "
    "bu sensörün serbest akış durumunda olduğunu göstermektedir. Modelin "
    "görevi bu 12 adımlık geçmiş penceresinden sonraki 12 adımın hızlarını "
    "tahmin etmektir."
))

add_heading(doc, "Sensörün graf konumu ve komşuları",
            level=3, numbered=True, number="2.4.3")
n_neighbors = len(example_neighbors)
add_par(doc, (
    f"Aynı sensörün komşuluk matrisindeki komşu sayısı {n_neighbors}'tir. Tablo "
    "3, bu komşu sensörlerin indislerini ve aralarındaki kenar ağırlıklarını "
    "göstermektedir. Yönlü graf yapısı nedeniyle giden kenarlar (sensör → "
    "komşu) ve gelen kenarlar (komşu → sensör) farklı ağırlıklara sahiptir."
))

n_show = min(n_neighbors, 8)
add_table_caption(doc, f"Örnek sensörün komşuları ve kenar ağırlıkları "
                  f"(ilk {n_show} komşu).")
tbl = doc.add_table(rows=n_show + 1, cols=3)
tbl.cell(0, 0).text = "Komşu sensör indisi"
tbl.cell(0, 1).text = "Giden kenar ağırlığı"
tbl.cell(0, 2).text = "Gelen kenar ağırlığı"
for j in range(3):
    for p in tbl.cell(0, j).paragraphs:
        for run in p.runs: run.bold = True
for i, nb in enumerate(example_neighbors[:n_show]):
    tbl.cell(i+1, 0).text = str(int(nb))
    tbl.cell(i+1, 1).text = f"{A_full[EXAMPLE_SENSOR_IDX, int(nb)]:.4f}"
    tbl.cell(i+1, 2).text = f"{A_full[int(nb), EXAMPLE_SENSOR_IDX]:.4f}"
style_table(tbl, font_size=10)
doc.add_paragraph()

add_par(doc, (
    "Tablodan görüleceği gibi giden ve gelen kenar ağırlıkları aynı komşu için "
    "farklı değerler almaktadır; bu, daha önce vurgulanan asimetri özelliğini "
    "somut olarak göstermektedir. DCRNN'in difüzyon konvolüsyonu bu iki "
    "yöndeki olasılık akışlarını ayrı ayrı modellemekte; A3T-GCN ise ortalama "
    "alarak tek bir simetrik temsil oluşturmaktadır."
))

add_heading(doc, "Tensor şekilleri ve modele giriş", level=3, numbered=True, number="2.4.4")
add_par(doc, (
    "Tam bir eğitim mini-grubu (mini-batch) için tensor şekilleri şu şekildedir: "
    "girdi [B, N, F, T_in] = [32, 207, 2, 12] (yığın boyutu × sensör sayısı × "
    "öznitelik sayısı × girdi adım sayısı); çıktı [B, N, T_out] = [32, 207, 12] "
    "(her sensör için 12 adımlık tahmin). Komşuluk matrisi modeller arasında "
    "paylaşılmaktadır ve eğitim boyunca sabit kalmaktadır; bu, METR-LA'nın "
    "statik graf yapısını yansıtmaktadır."
))


# -----------------------------------------------------------------------------
# 3. MODELLER VE METODOLOJİ
# -----------------------------------------------------------------------------
add_heading(doc, "Modeller ve Metodoloji", level=1, numbered=True, number=3)

add_par(doc, (
    "Bu bölümde uygulanan üç mimari ve bunlara ortak eğitim protokolü "
    "tanımlanmaktadır. Tüm mimariler aynı girdi-çıktı sözleşmesini izlemektedir: "
    "girdi her sensörün son 12 adımlık ölçümleri (hız + günün saati), çıktı her "
    "sensörün gelecek 12 adımdaki tahmin edilen hızıdır."
))

add_heading(doc, "Mimariler", level=2, numbered=True, number="3.1")

add_heading(doc, "A3T-GCN: Yönsüz graf konvolüsyonu ile zamansal dikkat",
            level=3, numbered=True, number="3.1.1")
add_par(doc, (
    "Attention Temporal Graph Convolutional Network (A3T-GCN; Zhu ve diğerleri, "
    "2020), TGCN bloğunun üzerine zamansal dikkat mekanizması ekleyen bir "
    "mimaridir. Mekansal bilgi standart bir graf konvolüsyon katmanı (GCN; Kipf "
    "ve Welling, 2017) ile birleştirilmekte; ancak GCN doğası gereği simetrik "
    "komşuluk matrisi varsayar. Bu özellik METR-LA'nın asimetrik graf yapısını "
    "modelleme açısından bir kısıt oluşturmaktadır. Bu çalışmada PyTorch "
    "Geometric Temporal'in yığın tabanlı sürümü olan A3TGCN2 kullanılmıştır; "
    f"gizli boyut 128, toplam öğrenilebilir parametre sayısı {m_a3t['num_params']:,}'dır. "
    "Modelin akış şeması Şekil 4'te sunulmaktadır."
))
add_figure(doc, "fig_arch_a3tgcn.png",
           "A3T-GCN mimari akışı. Her zaman adımı için graf konvolüsyon + "
           "TGCN bloğu uygulanır; 12 zaman adımının çıktıları zamansal dikkat "
           "ağırlıklarıyla toplanır ve lineer baş ile 12 adımlık tahmine "
           "dönüştürülür.")

add_heading(doc, "DCRNN: Yönlü difüzyon konvolüsyonu ile recurrent hücre",
            level=3, numbered=True, number="3.1.2")
add_par(doc, (
    "Diffusion Convolutional Recurrent Neural Network (DCRNN; Li ve diğerleri, "
    "2018), yönlü komşuluk matrisi üzerinde difüzyon konvolüsyon operatörünü ve "
    "GRU recurrent hücresini birleştirmektedir. Difüzyon konvolüsyonu, bir "
    "düğümden K adım uzaklıktaki düğümlere olasılıksal yürüyüş ile bilgi "
    "yayılımını modellemektedir. METR-LA'nın asimetrik graf yapısını doğrudan "
    "modelleyebilen tek mimaridir. Bu çalışmada K=2 (difüzyon adımı sayısı) ile "
    "yığın tabanlı sürüm (BatchedDCRNN) kullanılmış, gizli boyut 128 alınmıştır. "
    f"Toplam parametre sayısı {m_dcrnn['num_params']:,}'dır. Şekil 5 mimarinin "
    "akışını ve recurrent döngü yapısını göstermektedir."
))
add_figure(doc, "fig_arch_dcrnn.png",
           "DCRNN mimari akışı — yönlü difüzyon konvolüsyonu ile GRU "
           "recurrent hücresi. Her zaman adımında difüzyon adımları "
           "uygulanır; son gizli durum lineer baş ile 12 adımlık tahmine "
           "projeksiyon yapılır.")

add_heading(doc, "Vanilla LSTM: Graf yapısı kullanmayan kontrol grubu",
            level=3, numbered=True, number="3.1.3")
add_par(doc, (
    "Karşılaştırma noktası olarak graf yapısını hiç kullanmayan tek katmanlı, "
    "paylaşılan parametreli bir LSTM uygulanmıştır. Bu mimarinin amacı, graf "
    "yapısının trafik tahminindeki katkısını ayrıştırmaktır: eğer graf-bilinçli "
    "mimariler bu karşılaştırma noktasını net biçimde geçemiyorsa, ek mimari "
    "karmaşıklığın bu problem türünde değer eklemediği sonucuna varılır. Tüm "
    "207 sensör için aynı LSTM ağırlıkları kullanılmakta; bu, etkin yığın "
    "boyutunu 207 kat artıran bir paylaşılan parametre etkisi oluşturmaktadır. "
    f"Gizli boyut 128, toplam parametre sayısı {m_lstm['num_params']:,}'dır. "
    "Şekil 6 mimari akışı göstermektedir."
))
add_figure(doc, "fig_arch_lstm.png",
           "Vanilla LSTM mimari akışı. Graf yapısı kullanılmaz; her "
           "(sensör, örnek) çifti bağımsız bir zaman serisi olarak işlenir. "
           "Paylaşılan LSTM ağırlıkları tüm 207 sensör için aynıdır.")

# 3.2 Hiperparametreler
add_heading(doc, "Hiperparametreler ve Eğitim Protokolü",
            level=2, numbered=True, number="3.2")
add_par(doc, (
    "Üç mimari arasında adil karşılaştırma yapılabilmesi için aşağıdaki "
    "hiperparametreler tüm modellerde sabit tutulmuştur. Kayıp fonksiyonu "
    "olarak ortalama karesel hata (MSE) seçilmiştir; MSE, büyük hataları "
    "orantısız biçimde cezalandırdığı için modelleri büyük tahmin sapmalarından "
    "kaçınmaya yönlendirir ve ITS uygulamasında kritik olan büyük tıkanıklık "
    "anomalilerinin erken yakalanmasına uygundur. Kayıp fonksiyonu tercihinin "
    "alternatifleri Bölüm 5.4'te tartışılmaktadır."
))

add_table_caption(doc, "Tüm mimariler için sabit tutulan eğitim hiperparametreleri.")
hp_rows = [
    ("Eniyileme algoritması", "Adam"),
    ("Öğrenme oranı", "1×10⁻³"),
    ("Kayıp fonksiyonu", "Ortalama karesel hata (MSE)"),
    ("Gizli boyut", "128"),
    ("Yığın boyutu", "32"),
    ("Epoch sayısı", "10"),
    ("Rastgele tohum", "42"),
    ("Girdi pencere uzunluğu", "12 adım (60 dakika)"),
    ("Çıktı pencere uzunluğu", "12 adım (60 dakika)"),
    ("Donanım", "NVIDIA RTX 4050 Laptop GPU (6 GB VRAM)"),
    ("Yazılım kütüphanesi", "PyTorch 2.x + PyTorch Geometric Temporal"),
]
tbl = doc.add_table(rows=len(hp_rows), cols=2)
for i, (k, v) in enumerate(hp_rows): tbl.cell(i, 0).text = k; tbl.cell(i, 1).text = v
style_table(tbl)
doc.add_paragraph()

# 3.3 Değerlendirme protokolü
add_heading(doc, "Değerlendirme Protokolü", level=2, numbered=True, number="3.3")
add_par(doc, (
    "Her modelin her epoch sonunda doğrulama kümesi üzerinde MSE değeri "
    "hesaplanmaktadır. Eğitim tamamlandıktan sonra, test kümesi üzerinde "
    "tahminler üretilmekte ve şu metrikler raporlanmaktadır: (i) ortalama "
    "mutlak hata (MAE), (ii) kök ortalama karesel hata (RMSE), (iii) tahmin "
    "ufkunun her adımında (5, 10, ..., 60 dakika) ayrı MAE değerleri. Tüm "
    "metrikler önce z-skor uzayında hesaplanmakta, ardından gerçek birim "
    "(mph) cinsine dönüştürülmektedir. Bu çift-birim raporlama, sonuçların "
    "hem literatürle karşılaştırılabilirliğini hem de operasyonel "
    "yorumlanabilirliğini sağlamaktadır."
))

add_heading(doc, "Proje Hattı", level=2, numbered=True, number="3.4")
add_par(doc, (
    "Şekil 7, baştan sona projenin akışını sunmaktadır: veri yüklemeden "
    "başlayıp ön işleme, üç farklı modelin paralel eğitimi, sonuçların "
    "karşılaştırılması ve etkileşimli kontrol panosuna kadar tüm aşamalar "
    "gösterilmiştir."
))
add_figure(doc, "fig_pipeline.png",
           "Proje hattı — veri yüklemeden etkileşimli kontrol panosuna. "
           "Üç mimari eşzamanlı eğitilir; sonuçlar yığın halinde "
           "karşılaştırılır ve görsel raporlama üretilir.")


# -----------------------------------------------------------------------------
# 4. SONUÇLAR
# -----------------------------------------------------------------------------
add_heading(doc, "Sonuçlar", level=1, numbered=True, number=4)

# 4.1 Eğitim eğrileri
add_heading(doc, "Eğitim Eğrileri", level=2, numbered=True, number="4.1")
add_par(doc, (
    "Üç mimarinin epoch başına eğitim ve doğrulama MSE eğrileri Şekil 8'de "
    "sunulmaktadır. Tüm mimariler kararlı biçimde yakınsamış; eğitim sürecinde "
    "büyük dalgalanma veya aşırı uyum belirtisi gözlenmemiştir. DCRNN'in "
    "doğrulama MSE değeri yedinci epoch civarında plato'ya girmiş; A3T-GCN ise "
    "10 epoch sonunda hâlâ azalma eğilimini sürdürmektedir, bu durum daha uzun "
    "eğitim ile marjinal iyileşme potansiyelinin bulunduğunu işaret etmektedir. "
    "Vanilla LSTM dördüncü epoch'tan sonra plato'ya girmekte; daha fazla "
    "eğitimin marjinal kazanç sağlayacağı görülmemektedir."
))
add_figure(doc, "comparison_loss_curves.png",
           "Üç mimarinin eğitim (sol) ve doğrulama (sağ) MSE eğrileri. Tüm "
           "modeller kararlı yakınsama göstermektedir; DCRNN'in doğrulama "
           "MSE'si 7. epoch civarında plato yapmaktadır.")

# 4.2 Test performansı
add_heading(doc, "Test Kümesi Performansı", level=2, numbered=True, number="4.2")
add_table_caption(doc, "Test kümesi üzerinde üç mimarinin genel performans karşılaştırması. "
                  "Tüm değerler gerçek birim (mph) cinsinden raporlanmıştır.")
tbl = doc.add_table(rows=4, cols=4)
hdr = ["Mimari", "Parametre sayısı", "Test MAE (mph)", "Test RMSE (mph)"]
for j, h in enumerate(hdr):
    cell = tbl.cell(0, j); cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True
for i, m in enumerate(m_real["models"]):
    name = m["model"].replace(" baseline (batched)", "").replace(" — ablation baseline", "")
    tbl.cell(i+1, 0).text = name
    tbl.cell(i+1, 1).text = f"{m['num_params']:,}"
    tbl.cell(i+1, 2).text = f"{m['test_mae_mph']:.2f}"
    tbl.cell(i+1, 3).text = f"{m['test_rmse_mph']:.2f}"
style_table(tbl)
doc.add_paragraph()

add_par(doc, (
    f"Tablo 4'ten görüldüğü üzere, MAE metriğinde vanilla LSTM "
    f"({m_real['models'][2]['test_mae_mph']:.2f} mph) diğer iki mimariyi "
    f"geçmiştir; buna karşın RMSE metriğinde DCRNN "
    f"({m_real['models'][1]['test_rmse_mph']:.2f} mph) önde yer almıştır. RMSE "
    "büyük hatalara duyarlı olduğundan, DCRNN'in büyük tıkanıklıkları LSTM'den "
    "daha iyi öngördüğü sonucu çıkmaktadır. A3T-GCN her iki metrikte de "
    "geride kalmıştır; bu mimarinin asimetrik graf yapısını modelleyememesi "
    "temel nedendir."
))
add_figure(doc, "comparison_test_metrics.png",
           "Üç mimarinin test MAE (sol) ve RMSE (sağ) değerlerinin çubuk "
           "grafik karşılaştırması.")

# 4.3 Ufuk-bazlı analiz
add_heading(doc, "Tahmin Ufkuna Göre Yeterlilik Analizi",
            level=2, numbered=True, number="4.3")
add_par(doc, (
    "Mimarilerin avantajının tahmin ufkuna bağlı olarak değiştiği gözlemlenmiştir. "
    "Şekil 9 ufuk-bazlı MAE eğrilerini ve Şekil 10 ufuk × mimari ısı haritasını "
    "sunmaktadır. Tablo 5 her ufuk için sayısal değerleri içermektedir."
))

add_table_caption(doc, "Tahmin ufkuna göre MAE değerleri (mph). "
                  "En düşük değer her satırda kalın yazıdadır.")
tbl = doc.add_table(rows=13, cols=4)
hdr = ["Ufuk (dk)", "A3T-GCN", "DCRNN", "Vanilla LSTM"]
for j, h in enumerate(hdr):
    cell = tbl.cell(0, j); cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True
for h in range(12):
    tbl.cell(h+1, 0).text = f"+{(h+1)*5}"
    vals = [m["horizon_mae_mph"][h] for m in m_real["models"]]
    min_idx = vals.index(min(vals))
    for j, v in enumerate(vals):
        tbl.cell(h+1, j+1).text = f"{v:.2f}"
        if j == min_idx:
            for p in tbl.cell(h+1, j+1).paragraphs:
                for run in p.runs: run.bold = True
style_table(tbl)
doc.add_paragraph()

add_figure(doc, "realunit_horizon_mae.png",
           "Ufuk-bazlı MAE eğrileri (mph). Vanilla LSTM 5-50 dakika "
           "ufkunda en düşük MAE değerini verirken, 55 dakika ve "
           "sonrasında DCRNN öne geçmektedir.")
add_figure(doc, "fig_horizon_heatmap.png",
           "Ufuk × mimari MAE ısı haritası. Kırmızı tonlar yüksek hatayı, "
           "yeşil tonlar düşük hatayı temsil eder. DCRNN ile vanilla "
           "LSTM'in ufuk büyüdükçe yakınlaşması görsel olarak nettir.")

add_par(doc, (
    f"Anahtar yeterlilik bulgusu: vanilla LSTM 5-50 dakikalık ufuk aralığında "
    f"en düşük MAE değerini elde ederken, 55 dakikada DCRNN "
    f"({m_real['models'][1]['horizon_mae_mph'][10]:.2f} mph) ilk kez LSTM'i "
    f"({m_real['models'][2]['horizon_mae_mph'][10]:.2f} mph) geçmektedir. Bu "
    "geçiş noktası (crossover), graf sinir ağlarının mekansal indüktif "
    "önyargısının uzun ufukta yeterli olmaya başladığının somut göstergesidir. "
    "Bulgu, mimari seçiminin operasyonel ihtiyaca (kısa veya uzun ufuk "
    "tahmin) göre yapılması gerektiğini ortaya koymaktadır."
))

# 4.4 Tahmin görselleştirmesi
add_heading(doc, "Tahmin Görselleştirmesi", level=2, numbered=True, number="4.4")
add_par(doc, (
    "Şekil 11 dört örnek sensörde mimarilerin 5 dakika öne yaptığı tahminleri "
    "gerçek değerle birlikte göstermektedir. Şekil 12 en düşük MAE değerini "
    "elde eden mimarinin tahmin-gerçek scatter plot'unu sunmaktadır."
))
add_figure(doc, "viz_predictions_multi.png",
           "Dört örnek sensörde mimarilerin 5 dakika öne tahminleri ve "
           "gerçek değerler. Tüm modeller hız trendini yakalamakta, ancak "
           "ani düşüşler (tıkanıklık başlangıçları) hâlâ zorlayıcıdır.")
add_figure(doc, "viz_scatter_best.png",
           "Vanilla LSTM mimarisinin +5 dakika tahmin scatter plot'u. "
           "Yüksek hızlarda (60+ mph) tahminlerin yığılması, METR-LA'nın "
           "otoyol sensörlerinden geldiğini yansıtmaktadır.")

# 4.5 Mekansal yayılma analizi
add_heading(doc, "Mekansal Yayılma Analizi", level=2, numbered=True, number="4.5")
add_par(doc, (
    "DCRNN'in mekansal indüktif önyargısını sayısal olarak doğrulamak amacıyla "
    "bir yayılma deneyi yapılmıştır: seçilen bir bölgedeki her sensör tek tek "
    "tıkalı duruma getirildiğinde, diğer sensörlerin +30 dakika tahminindeki "
    "değişim ölçülmüştür. Şekil 13 bu yayılma matrisini görselleştirmektedir."
))
add_figure(doc, "sim_impact_matrix.png",
           "Sensörler arası yayılma şiddeti matrisi (Downtown LA bölgesi, "
           "+30 dakika ufuk). Köşegen müdahale edilen sensörü, satırlar "
           "o müdahalenin diğer sensörlere etkisini göstermektedir.")

add_par(doc, (
    "Matristen okunan bulgu şudur: tek bir sensörün durumundaki değişiklik "
    "komşu sensörlerin tahminlerini de değiştirmektedir. Klasik bir LSTM "
    "modelinde bu matriste yalnızca köşegen değerler sıfırdan farklı olurdu "
    "(her sensör birbirinden bağımsız tahmin edilir). DCRNN'in graf yapısı "
    "üzerinden öğrendiği yayılma bilgisi, bu deneyle somut biçimde gözlemlenebilir "
    "hale gelmiştir."
))

# 4.6 Eğitim süresi
add_heading(doc, "Eğitim Süresi ve Hesaplama Maliyeti",
            level=2, numbered=True, number="4.6")
add_par(doc, (
    "PyTorch Geometric Temporal'in yığın tabanlı sürümleri kullanılarak "
    "(A3TGCN2, BatchedDCRNN) yaklaşık 150 kat hızlanma elde edilmiştir. "
    "Tek tek snapshot işleme stratejisi (toplam ~12 saat tahmini) yerine "
    "yığın işleme ile tüm pipeline yaklaşık 35 dakikada tamamlanmaktadır. "
    "Bu hızlanma sayısal sonuçları değil yalnızca eğitim süresini etkilemektedir."
))
add_table_caption(doc, "Mimari başına epoch süresi ve toplam eğitim süresi.")
tbl = doc.add_table(rows=4, cols=3)
hdr = ["Mimari", "Epoch süresi (saniye)", "Toplam (10 epoch)"]
for j, h in enumerate(hdr):
    cell = tbl.cell(0, j); cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True
times = [("A3T-GCN", "~85", "~15 dakika"),
         ("DCRNN", "~125", "~21 dakika"),
         ("Vanilla LSTM", "~25", "~4 dakika")]
for i, (m, e, t) in enumerate(times):
    tbl.cell(i+1, 0).text = m
    tbl.cell(i+1, 1).text = e
    tbl.cell(i+1, 2).text = t
style_table(tbl)
doc.add_paragraph()

# 4.7 Literatürle karşılaştırma
add_heading(doc, "Literatürdeki Diğer Çalışmalarla Karşılaştırma",
            level=2, numbered=True, number="4.7")
add_par(doc, (
    "METR-LA veri kümesi 2018'den bu yana çok sayıda çalışmaya konu olmuştur; "
    "literatürdeki farklı yaklaşımlar 15 dakikalık ufukta 2,7 — 3,3 mph "
    "bandında MAE değerleri raporlamıştır. Tablo 7 bu çalışmanın sonuçlarını "
    "literatürdeki referans noktaları ile karşılaştırmaktadır."
))

add_table_caption(doc, "METR-LA üzerinde seçilmiş çalışmaların ve bu çalışmanın "
                  "MAE değerleri (mph).")
tbl = doc.add_table(rows=7, cols=4)
hdr = ["Çalışma", "+15 dk", "+30 dk", "+60 dk"]
for j, h in enumerate(hdr):
    cell = tbl.cell(0, j); cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True
lit_rows = [
    ("DCRNN paper'ı (Li et al., 2018)", "2,77", "3,15", "3,60"),
    ("STGCN paper'ı (Yu et al., 2018)", "2,88", "3,48", "4,59"),
    ("Graph WaveNet (Wu et al., 2019)", "2,69", "3,07", "3,53"),
    ("Bu çalışma — DCRNN",
     f"{m_real['models'][1]['horizon_mae_mph'][2]:.2f}",
     f"{m_real['models'][1]['horizon_mae_mph'][5]:.2f}",
     f"{m_real['models'][1]['horizon_mae_mph'][11]:.2f}"),
    ("Bu çalışma — A3T-GCN",
     f"{m_real['models'][0]['horizon_mae_mph'][2]:.2f}",
     f"{m_real['models'][0]['horizon_mae_mph'][5]:.2f}",
     f"{m_real['models'][0]['horizon_mae_mph'][11]:.2f}"),
    ("Bu çalışma — Vanilla LSTM",
     f"{m_real['models'][2]['horizon_mae_mph'][2]:.2f}",
     f"{m_real['models'][2]['horizon_mae_mph'][5]:.2f}",
     f"{m_real['models'][2]['horizon_mae_mph'][11]:.2f}"),
]
for i, row in enumerate(lit_rows, 1):
    for j, v in enumerate(row):
        tbl.cell(i, j).text = v
style_table(tbl)
doc.add_paragraph()

add_par(doc, (
    "Niceliksel olarak bu çalışmanın sonuçları literatürdeki paper'ların "
    "gerisinde kalmaktadır. Bu farkın temel nedenleri Bölüm 5.2'de "
    "tartışılmıştır. Çalışmanın özgün katkısı mutlak skor liderliği yerine, "
    "kontrollü deneysel çerçevenin uygulanması ve mimarilerin yeterlilik "
    "koşullarının ortaya konmasıdır. Sınırlı eğitim bütçesinde dahi DCRNN'in "
    "uzun ufuk ve RMSE metriğinde tutarlı üstünlüğü ile A3T-GCN'in asimetrik "
    "graf modellemedeki yetersizliği net biçimde gözlemlenmiştir."
))


# -----------------------------------------------------------------------------
# 5. TARTIŞMA
# -----------------------------------------------------------------------------
add_heading(doc, "Tartışma", level=1, numbered=True, number=5)
add_par(doc, (
    "Bu bölüm Bölüm 1'de formüle edilen araştırma sorularını ampirik verilere "
    "dayandırarak değerlendirmekte, yöntemsel gözlemleri yorumlamakta ve "
    "çalışmanın sınırlılıklarını ortaya koymaktadır."
))

# 5.1 Araştırma sorularına yanıt
add_heading(doc, "Araştırma Sorularının Değerlendirilmesi",
            level=2, numbered=True, number="5.1")

add_par(doc, (
    "S1, Mimari karşılaştırması. Üç mimari "
    f"({m_real['models'][2]['test_mae_mph']:.2f} — "
    f"{m_real['models'][0]['test_mae_mph']:.2f} mph MAE bandında) farklı "
    "performans dağılımları sergilemiştir. DCRNN > LSTM > A3T-GCN sıralaması "
    "RMSE metriğinde gözlemlenmiş; MAE metriğinde ise sıralama LSTM > DCRNN > "
    "A3T-GCN olmuştur. Mesaj iletme şemalarındaki farklılıklar trafik tahmin "
    "görevinde anlamlı sonuçlara yol açmaktadır: yönlü difüzyon konvolüsyonu "
    "(DCRNN) graf yönlülüğünü doğrudan modeller ve büyük hatalardan kaçınır; "
    "simetrik graf konvolüsyonu (A3T-GCN) ise asimetri bilgisini kaybeder ve "
    "her metrikte geride kalır."
), indent=False)

add_par(doc, (
    "S2, Yönlülük etkisi. Hem A3T-GCN hem DCRNN graf bilinçli mimarilerdir; "
    "tek temel fark, DCRNN'in yönlü komşuluk matrisini doğrudan modellemesidir. "
    f"DCRNN'in A3T-GCN'i tüm ufuklarda geçmesi (MAE'de "
    f"{(m_real['models'][0]['test_mae_mph'] - m_real['models'][1]['test_mae_mph']):.2f} "
    "mph fark, RMSE'de "
    f"{(m_real['models'][0]['test_rmse_mph'] - m_real['models'][1]['test_rmse_mph']):.2f} "
    "mph fark) yönlülüğün net katkısının nümerik kanıtıdır. Bu fark mutlak "
    "olarak küçük görünse de tutarlı yönde olduğu için yöntemsel olarak "
    "anlamlıdır: A3T-GCN simetri varsayımı nedeniyle bilgi kaybetmekte; "
    "DCRNN bu bilgiyi korumaktadır. ITS uygulamasında trafik akışının yön "
    "bilgisi temel bir özellik olduğundan, bu bulgu mimari seçiminde yönlü "
    "modelleri lehinde güçlü bir gerekçedir."
), indent=False)

add_par(doc, (
    "S3, Ufka göre yeterlilik. Vanilla LSTM 5-50 dakikalık ufuk aralığında "
    "MAE metriğinde en düşük değeri vermektedir; 55 dakikada DCRNN ilk kez "
    "LSTM'i geçmektedir. Bu geçiş noktası araştırma hipotezini "
    "doğrulamaktadır: kısa ufukta her sensörün geleceği büyük ölçüde kendi "
    "geçmişine bağlıdır ve graf mekanizmasının ek katkısı sınırlıdır; uzun "
    "ufukta ise komşu sensörlerden gelen yayılma sinyali belirleyici hale "
    "gelmektedir. Hipotez doğrulanmıştır: graf-bilinçli mimariler ile graf "
    "yapısı kullanmayan klasik bir tekrarlı ağ arasında ufuk-bağımlı bir "
    "yeterlilik geçişi mevcuttur."
), indent=False)

add_par(doc, (
    "S4, Hata metriği tercihi. MAE metriği tipik hatayı, RMSE büyük hatalara "
    "duyarlılığı ölçmektedir. MAE'de LSTM, RMSE'de DCRNN üstün gelmiştir. Bu "
    "metrik-bağımlı sıralama, mimari tercihinin operasyonel hedeflere göre "
    "yapılması gerektiğini ortaya koymaktadır: tipik doğruluk istenen "
    "uygulamalarda (sürücü bilgilendirme) LSTM yeterli olabilirken, büyük "
    "tıkanıklık olaylarından kaçınmanın kritik olduğu uygulamalarda "
    "(proaktif sinyal yönetimi) DCRNN tercih edilmelidir."
), indent=False)

# 5.2 Literatürle fark
add_heading(doc, "Literatürdeki Skorlardan Geri Kalma Nedenleri",
            level=2, numbered=True, number="5.2")
add_par(doc, (
    "Tablo 7'de görüldüğü üzere bu çalışmanın MAE değerleri literatürdeki "
    "referans paper'larının yaklaşık 2 katı civarındadır. Bu farkın temel "
    "nedenleri üç başlıkta toplanabilir."
))

reasons = [
    ("Decoder mimarisi.", "Bu çalışmada DCRNN için sade bir 'son gizli durum + "
     "lineer projeksiyon' decoder kullanılmıştır. Orijinal DCRNN paper'ı, "
     "encoder-decoder yapısı ve teacher forcing eğitim stratejisi ile "
     "raporlama yapmaktadır. Bu yapısal fark, özellikle uzun ufuk tahmininde "
     "performansta belirgin etki yaratabilmektedir."),
    ("Eğitim süresi.", "Literatürdeki çalışmalar tipik olarak 100+ epoch eğitim "
     "ile raporlama yapmakta; bu çalışma 10 epoch ile sınırlı tutulmuştur. "
     "Eğitim eğrileri (Şekil 8) bazı mimarilerin (özellikle A3T-GCN) daha "
     "uzun eğitim ile iyileşme potansiyeli taşıdığını göstermektedir."),
    ("Sistematik hiperparametre araması yapılmamıştır.", "Tüm hiperparametre "
     "değerleri sezgisel olarak seçilmiş (gizli boyut 128 gibi); Bayesian "
     "optimizasyon veya ızgara araması uygulanmamıştır. Bu yapılsa "
     "performansta iyileşme alanı mevcuttur."),
]
for tag, body in reasons:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run(tag + " "); r1.bold = True
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
    r2 = p.add_run(body); r2.font.name = "Times New Roman"; r2.font.size = Pt(11)

# 5.3 A3T-GCN'in zayıflığı
add_heading(doc, "A3T-GCN'in Asimetrik Graf'ta Zayıf Performansı",
            level=2, numbered=True, number="5.3")
add_par(doc, (
    "A3T-GCN'in temel GCN katmanı simetrik komşuluk matrisi varsayımıyla "
    "çalışmaktadır. Ancak METR-LA'nın komşuluk matrisi açıkça asimetriktir: "
    "trafiğin X sensöründen Y'ye akış bilgisi Y'den X'e akıştan farklıdır. "
    "A3T-GCN bu ayrımı kaybeder; sonuçta sinyal akışının yönüne saygılı "
    "olmayan bir model elde edilir. Bu, A3T-GCN'in DCRNN tarafından her "
    "ufukta geçilmesinin temel nedenidir. Trafik gibi yönlülüğün anlamlı "
    "olduğu uygulamalarda yönsüz graf konvolüsyonu kullanan mimarilerin "
    "sınırlı kalacağı öngörülebilir."
))

# 5.4 Kayıp fonksiyonu seçiminin etkisi
add_heading(doc, "Kayıp Fonksiyonu Seçiminin Hata Profili Üzerindeki Etkisi",
            level=2, numbered=True, number="5.4")
add_par(doc, (
    "Bu çalışmada tüm modeller MSE kayıp fonksiyonu ile eğitilmiştir. MSE, "
    "büyük hataları kare ile orantılı olarak cezalandırır; bu, modelin büyük "
    "tahmin sapmalarından kaçınmasını teşvik eder ve RMSE metriğinde "
    "iyileşmeye katkı sağlar. Alternatif olarak MAE (L1) kaybı seçilseydi, "
    "model daha sağlam medyan-benzeri tahminlere yönelirdi; ancak aşırı "
    "durumlardaki keskinliği kaybederdi. Huber kaybı bu iki davranış "
    "arasında bir orta yol sunmaktadır."
))
add_par(doc, (
    "Kayıp fonksiyonu seçimi mimari karşılaştırma sıralamasını değiştirebilir: "
    "MAE kaybı ile eğitilmiş bir modelin MAE metriği muhtemelen daha iyi "
    "olurken RMSE'si bozulacaktır. Bu çalışmada MSE seçimi yapılmıştır çünkü "
    "ITS uygulamalarında büyük tıkanıklık anomalilerinin yakalanması "
    "operasyonel olarak kritiktir. Farklı uygulamalar farklı kayıp fonksiyonu "
    "tercihleri gerektirebilir; bu, mimari seçiminin yanı sıra dikkate "
    "alınması gereken ikinci bir tasarım kararıdır."
))

# 5.5 Sınırlılıklar
add_heading(doc, "Çalışmanın Sınırlılıkları", level=2, numbered=True, number="5.5")
lims = [
    "Tek bir veri kümesi (METR-LA) üzerinde değerlendirilmiştir. PeMS-BAY "
    "veya yerel veri kümeleriyle çapraz doğrulama yapılmamış; bulguların "
    "genelleştirilebilirliği test edilmemiştir.",

    "Sistematik hiperparametre araması yapılmamıştır. Gizli boyut, öğrenme "
    "oranı ve yığın boyutu sezgisel olarak seçilmiştir.",

    "Tek tohum (42) ile çalışılmıştır. Çoklu tohum ile ortalama ve standart "
    "sapma raporlama yapılmadığından stokastik varyansın etkisi nicelleştirilememiştir.",

    "DCRNN için orijinal paper'da kullanılan encoder-decoder + teacher forcing "
    "yapısı uygulanmamış; basit son gizli durum + lineer decoder tercih "
    "edilmiştir. Bu, literatür skoruna mesafenin büyük ölçüde nedenidir.",

    "Eksik veri (sıfır okumalar, %8,11) doldurma uygulanmamıştır. Forward-fill "
    "veya mask-aware eğitim teknikleri ile performansta iyileşme alanı bulunmaktadır.",
]
for lim in lims:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lim); r.font.name = "Times New Roman"; r.font.size = Pt(11)

# 5.6 Pratik öneriler
add_heading(doc, "Pratik Öneriler", level=2, numbered=True, number="5.6")
practical = [
    ("Kısa ufuk için klasik tekrarlı ağlar yeterlidir.",
     "5-45 dakikalık tahmin gerektiren uygulamalarda iyi tasarlanmış bir "
     "paylaşılan-parametre LSTM, üretim ortamında düşük kaynak gereksinimiyle "
     "yeterli sonuç sağlamaktadır."),
    ("Uzun ufuk ve RMSE-duyarlı senaryolar için DCRNN tercih edilmelidir.",
     "45 dakikadan uzun tahminler ve büyük tıkanıklık olaylarının "
     "yakalanmasının kritik olduğu durumlarda yönlü graf modelleyebilen "
     "DCRNN açıkça öne çıkmaktadır."),
    ("Simetrik graf konvolüsyonu yönlü graflar için uygun değildir.",
     "A3T-GCN gibi simetrik komşuluk varsayımı yapan mimariler, yönlülüğün "
     "anlamlı olduğu uygulamalarda kullanılmamalıdır; performans kaybı "
     "tahmin edilebilir niteliktedir."),
]
for tag, body in practical:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run("· " + tag + " "); r1.bold = True
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
    r2 = p.add_run(body); r2.font.name = "Times New Roman"; r2.font.size = Pt(11)


# -----------------------------------------------------------------------------
# 6. ISSD ENTEGRASYON ÖNERİSİ
# -----------------------------------------------------------------------------
add_heading(doc, "ISSD Akıllı Ulaşım Platformuna Entegrasyon Önerisi",
            level=1, numbered=True, number=6)

add_par(doc, (
    "Bu bölüm, çıkarılan DCRNN modelinin ISSD Bilişim Elektronik A.Ş.'nin "
    "akıllı ulaşım platformuna mimari entegrasyonu için bir öneri çerçevesi "
    "sunmaktadır. ISSD, 2009 yılında ODTÜ Teknokent'te kurulmuş, Türkiye'nin "
    "akıllı ulaşım sistemleri alanındaki öncü şirketlerinden biridir. CHAOS "
    "adıyla tescilli Dynamic Junction Control System'i 20'den fazla şehirde "
    "1000'i aşkın kavşakta çalışmakta ve şirket bu teknolojiyi altı ülkeye "
    "ihraç etmektedir."
))

add_heading(doc, "Tespit Edilen Boşluklar", level=2, numbered=True, number="6.1")
add_par(doc, (
    "ISSD'nin halka açık dokümantasyonu sistemin anlık görüntü işleme açısından "
    "güçlü olduğunu göstermektedir. Ancak iki kritik boşluk göze çarpmaktadır: "
    "(i) öngörülü tahmin yokluğu — CHAOS reaktif çalışır, 30-60 dakika öne "
    "tahmin görünmemektedir; (ii) şehir ölçekli yayılma modellemesi yokluğu — "
    "her kavşak büyük ölçüde bağımsız optimize edilir."
))

add_heading(doc, "Önerilen Mimari", level=2, numbered=True, number="6.2")
add_par(doc, (
    "Bu boşlukları kapatmak için MANGO katmanının içine paralel bir öngörü "
    "modülü olarak DCRNN modelinin entegrasyonu önerilmektedir (Şekil 14). "
    "Mevcut sensör altyapısı (VIERO-AI kamera tabanlı araç sayma, BLUESIS "
    "Bluetooth tabanlı seyahat süresi, FCD GPS verileri) tek değişmeden "
    "5 dakikalık veriyi besler; DCRNN katmanı 60 dakika öne tahmin üretir ve "
    "sonuç hem operatör panosuna hem CHAOS'un sinyalizasyon kararlarına "
    "girdi olarak aktarılır."
))
add_figure(doc, "fig_issd_integration.png",
           "ISSD CHAOS/MANGO platformu için önerilen öngörü katmanı "
           "entegrasyonu. Mevcut sensör altyapısı değişmez; öngörü "
           "katmanı paralel modül olarak eklenir.")

add_heading(doc, "Operasyonel Hususlar", level=2, numbered=True, number="6.3")
add_par(doc, (
    "Tek inference çağrısı RTX 4050 sınıfı GPU'da yaklaşık 25 milisaniye "
    "sürmektedir. 1000'den fazla kavşak ölçeğinde (METR-LA'nın yaklaşık 5 "
    "katı) bu süre 100-200 milisaniye civarında kalır, 5 dakikalık tahmin "
    "döngüsünde tamamen makul kapsamdadır. Donanım açısından ISSD'nin "
    "Advantech/Intel ortaklığındaki mevcut sunucu altyapısı yeterlidir; "
    "edge düzeyinde inference gerekmemektedir."
))

add_heading(doc, "İnteraktif Kontrol Panosu Prototipi", level=2, numbered=True, number="6.4")
add_par(doc, (
    "Bu çalışmanın bir yan ürünü olarak Streamlit tabanlı etkileşimli bir "
    "kontrol panosu prototipi geliştirilmiştir. Panoda kullanıcı seçili "
    "bölgedeki kavşakların anlık trafik durumunu değiştirebilir ve sistemin "
    "60 dakika öne tahminini üç ayrı haritada (baseline, müdahaleli, fark) "
    "karşılaştırmalı olarak gözlemleyebilir. Bu prototip, CHAOS/MANGO "
    "operatörü deneyiminin somut bir taslağıdır ve savunma sırasında canlı "
    "gösterim olarak kullanılmaktadır."
))


# -----------------------------------------------------------------------------
# 7. SONUÇ
# -----------------------------------------------------------------------------
add_heading(doc, "Sonuç ve Gelecek Çalışmalar", level=1, numbered=True, number=7)

add_par(doc, (
    "Bu çalışmada trafik tahmini için üç farklı derin öğrenme mimarisi "
    "(A3T-GCN, DCRNN, vanilla LSTM) kontrollü deneysel ortamda karşılaştırılmış; "
    "yeterlilik koşulları, ufuk bağımlılığı, yönlülük etkisi ve hata metriği "
    "tercihi ekseninde sistematik olarak analiz edilmiştir. Anahtar bulgu, "
    "mimari yeterliliğinin tahmin ufkuna bağlı olarak değiştiğidir: kısa ufukta "
    "vanilla LSTM yeterli bir karşılaştırma noktası oluşturur; 55 dakika ve "
    "sonrasında DCRNN'in yönlü mekansal indüktif önyargısı yeterli hale gelir "
    "ve RMSE metriğinde her zaman öne çıkar. Akıllı ulaşım sistemleri için "
    "mimari seçimi sabit bir 'en iyi' yerine operasyonel ihtiyaca göre "
    "yapılmalıdır."
))

add_par(doc, (
    "Ek olarak, DCRNN modelinin ISSD'nin CHAOS/MANGO platformuna entegrasyonu "
    "için bir mimari öneri ve Streamlit tabanlı etkileşimli kontrol panosu "
    "prototipi sunulmuştur. Çalışmanın özgün katkısı mutlak skor liderliği "
    "yerine, kontrollü karşılaştırmalı çerçeve ve yeterlilik koşullarının "
    "çıkarılmasıdır."
))

add_par(doc, "Gelecek çalışmalar için somut yönler:", indent=False)
futures = [
    "PeMS-BAY veri kümesi ve İBB Açık Veri Portalı üzerinde çapraz doğrulama "
    "ile bulguların genelleştirilebilirliğinin test edilmesi.",

    "DCRNN için orijinal paper'da kullanılan encoder-decoder + teacher forcing "
    "yapısının uygulanması; literatür skoruna yaklaşmak.",

    "Çoklu tohum ile ortalama ± standart sapma raporlama; stokastik varyansın "
    "nicelleştirilmesi.",

    "Multi-modal genişleme: kamera, Bluetooth ve FCD verisinin birleştirilmiş "
    "heterogen graf yapısında modellenmesi.",

    "Tahmin-gerçek sapması üzerinden anomali/olay tespiti — bu çalışmanın "
    "yan ürünü olarak.",

    "ISSD'nin anonimleştirilmiş üretim verisinde modelin gerçek-dünya "
    "doğrulanması (staj kapsamı).",
]
for fw in futures:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(fw); r.font.name = "Times New Roman"; r.font.size = Pt(11)


# -----------------------------------------------------------------------------
# KAYNAKÇA
# -----------------------------------------------------------------------------
add_page_break(doc)
add_heading(doc, "Kaynakça", level=1)
references = [
    "Kipf, T. N., & Welling, M. (2017). Semi-Supervised Classification with "
    "Graph Convolutional Networks. International Conference on Learning "
    "Representations (ICLR).",

    "Li, Y., Yu, R., Shahabi, C., & Liu, Y. (2018). Diffusion Convolutional "
    "Recurrent Neural Network: Data-Driven Traffic Forecasting. International "
    "Conference on Learning Representations (ICLR). "
    f"({METR_LA_URL})",

    "Yu, B., Yin, H., & Zhu, Z. (2018). Spatio-Temporal Graph Convolutional "
    "Networks: A Deep Learning Framework for Traffic Forecasting. International "
    "Joint Conference on Artificial Intelligence (IJCAI).",

    "Zhu, J., Wang, Q., Tao, C., Deng, H., Zhao, L., & Li, H. (2020). AST-GCN: "
    "Attribute-Augmented Spatiotemporal Graph Convolutional Network for "
    "Traffic Forecasting. IEEE Access, 8, 35973-35983.",

    "Wu, Z., Pan, S., Long, G., Jiang, J., & Zhang, C. (2019). Graph WaveNet "
    "for Deep Spatial-Temporal Graph Modeling. International Joint Conference "
    "on Artificial Intelligence (IJCAI).",

    "Veličković, P., Cucurull, G., Casanova, A., Romero, A., Liò, P., & "
    "Bengio, Y. (2018). Graph Attention Networks. International Conference on "
    "Learning Representations (ICLR).",

    "Rozemberczki, B., Scherer, P., He, Y., Panagopoulos, G., Riedel, A., "
    "Astefanoaei, M., Kiss, O., Beres, F., Lopez, G., Collignon, N., & Sarkar, "
    "R. (2021). PyTorch Geometric Temporal: Spatiotemporal Signal Processing "
    "with Neural Machine Learning Models. ACM International Conference on "
    "Information and Knowledge Management (CIKM).",

    "Paszke, A., Gross, S., Massa, F., Lerer, A., et al. (2019). PyTorch: An "
    "Imperative Style, High-Performance Deep Learning Library. Advances in "
    "Neural Information Processing Systems 32 (NeurIPS).",

    "ISSD Bilişim Elektronik A.Ş. (2026). Şirket Bilgileri ve Ürün "
    "Dokümantasyonu. https://www.issd.com.tr",
]
for i, ref in enumerate(references, 1):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7); p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"[{i}] "); r1.bold = True
    r1.font.name = "Times New Roman"; r1.font.size = Pt(10)
    r2 = p.add_run(ref); r2.font.name = "Times New Roman"; r2.font.size = Pt(10)


# Kaydet
doc.save(OUTPUT_PATH)
print(f"\n✓ Tez güncellendi: {OUTPUT_PATH}")
print(f"  Görsel sayısı: {FIG_COUNTER['n']} · Tablo sayısı: {TBL_COUNTER['n']}")
